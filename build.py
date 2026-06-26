#!/usr/bin/env python3
"""build.py – Flexible parser + adaptive data injector for pedagogy-hub.html"""

import json, re, base64, os
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

SCRIPT_DIR = Path(os.path.abspath(__file__)).parent if '__file__' in dir() else Path('/home/claude')

UNICODE_MAP = {
    '\u2018':"'",'\u2019':"'",'\u201c':'"','\u201d':'"',
    '\u2013':'-','\u2014':'--','\u2026':'...','\u00a0':' ',
    '\uf0b7':'','\uf0a7':'','\uf020':'',
}
def clean(text):
    if not text: return ''
    for src,dst in UNICODE_MAP.items(): text=text.replace(src,dst)
    text=re.sub(r'[\ue000-\uf8ff]','',text)
    text=re.sub(r'\s+',' ',text)
    return text.strip()

def para_text(p):
    return clean(''.join(run.text or '' for run in p.runs))

def para_text_preserve_newlines(p):
    """Return paragraph text with internal line breaks preserved."""
    return clean(''.join(run.text or '' for run in p.runs))

def is_tech(p):
    return bool(re.match(r'^\(Tech\)', para_text(p), re.IGNORECASE))

def strip_tech(t):
    return re.sub(r'^\(Tech\)\s*','',t,flags=re.IGNORECASE).strip()

def para_style(p):
    return (p.style.name or '').strip()

def is_heading(p):
    return any(s in para_style(p) for s in ('Title','Heading'))

def has_image(p):
    return bool(p._element.findall('.//' + qn('a:blip')))

def compress_image(image_bytes, max_width=900, quality=72):
    """Compress image to JPEG if over threshold, preserving aspect ratio."""
    try:
        from PIL import Image
        import io
        img = Image.open(io.BytesIO(image_bytes))
        # Convert RGBA/P to RGB for JPEG
        if img.mode in ('RGBA', 'P', 'LA'):
            bg = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            bg.paste(img, mask=img.split()[-1] if img.mode in ('RGBA','LA') else None)
            img = bg
        elif img.mode != 'RGB':
            img = img.convert('RGB')
        # Resize if too wide
        if img.width > max_width:
            ratio = max_width / img.width
            img = img.resize((max_width, int(img.height * ratio)), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format='JPEG', quality=quality, optimize=True)
        return buf.getvalue(), 'image/jpeg'
    except Exception:
        return image_bytes, ('image/png' if image_bytes[:4]==b'\x89PNG' else 'image/jpeg')

def extract_image_b64(p, docx_path):
    blips = p._element.findall('.//' + qn('a:blip'))
    if not blips: return None
    rId = blips[0].get(qn('r:embed'))
    if not rId: return None
    doc2 = Document(docx_path)
    for rel in doc2.part.rels.values():
        if rel.reltype.endswith('/image') and rel.rId == rId:
            image_bytes = rel.target_part.blob
            # Compress if over 100KB
            if len(image_bytes) > 100 * 1024:
                image_bytes, mime = compress_image(image_bytes)
            else:
                mime = 'image/png' if image_bytes[:4]==b'\x89PNG' else 'image/jpeg'
            b64 = base64.b64encode(image_bytes).decode('utf-8')
            return f'data:{mime};base64,{b64}'
    return None

# ── numbering lookup ───────────────────────────────────────────────────────
def build_numbering_lookup(doc):
    numbering_part = None
    for rel in doc.part.rels.values():
        if 'numbering' in rel.reltype:
            numbering_part = rel.target_part
            break
    if not numbering_part: return {}, {}
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    root = etree.fromstring(numbering_part.blob)
    abstract = {}
    for abNum in root.findall(f'{{{ns}}}abstractNum'):
        abId = abNum.get(f'{{{ns}}}abstractNumId')
        abstract[abId] = {}
        for lvl in abNum.findall(f'{{{ns}}}lvl'):
            ilvl = lvl.get(f'{{{ns}}}ilvl')
            lvlText = lvl.find(f'{{{ns}}}lvlText')
            char = lvlText.get(f'{{{ns}}}val','') if lvlText is not None else ''
            abstract[abId][ilvl] = char
    num_to_abstract = {}
    for num in root.findall(f'{{{ns}}}num'):
        numId = num.get(f'{{{ns}}}numId')
        abRef = num.find(f'{{{ns}}}abstractNumId')
        if abRef is not None:
            num_to_abstract[numId] = abRef.get(f'{{{ns}}}val')
    return abstract, num_to_abstract

def get_bullet_level(p, abstract, num_to_abstract):
    numPr = p._element.find('.//' + qn('w:numPr'))
    ilvl_val = 0; numId_val = None
    if numPr is not None:
        ilvl_el = numPr.find(qn('w:ilvl'))
        numId_el = numPr.find(qn('w:numId'))
        if ilvl_el is not None: ilvl_val = int(ilvl_el.get(qn('w:val'),'0'))
        if numId_el is not None: numId_val = numId_el.get(qn('w:val'))
    if ilvl_val >= 1: return 1
    if numId_val and numId_val in num_to_abstract:
        abId = num_to_abstract[numId_val]
        char = abstract.get(abId,{}).get(str(ilvl_val),'')
        if char == 'o': return 1
    pPr = p._element.find(qn('w:pPr'))
    if pPr is not None:
        ind = pPr.find(qn('w:ind'))
        if ind is not None:
            ind_left = ind.get(qn('w:left'))
            if ind_left and int(ind_left) >= 400: return 1
    return 0

def cell_bullets(cell, abstract, num_to_abstract):
    bullets = []
    for p in cell.paragraphs:
        t = para_text(p)
        if not t: continue
        tech = is_tech(p)
        level = get_bullet_level(p, abstract, num_to_abstract)
        bullets.append({'text': strip_tech(t) if tech else t, 'tech': tech, 'level': level})
    return bullets

def cell_plain(cell):
    return clean(cell.text)

# ── aims cell parser ───────────────────────────────────────────────────────
TAG_RE = re.compile(r'^\[([^\]]+)\]')

def parse_aims_cell(cell):
    paras = [clean(p.text) for p in cell.paragraphs]
    paras = [p for p in paras if p]
    blocks = []; current_tag = None; current_lines = []
    def flush():
        if current_tag is not None:
            blocks.append({'tag': current_tag, 'text': ' '.join(current_lines).strip()})
    for para in paras:
        m = TAG_RE.match(para)
        if m:
            flush(); current_tag = m.group(1).strip()
            rest = para[m.end():].strip()
            current_lines = [rest] if rest else []
        else:
            if current_tag is not None: current_lines.append(para)
            else: blocks.append({'tag': None, 'text': para})
    flush()
    return blocks if blocks else None

# ── table identity ─────────────────────────────────────────────────────────
def identify_table(tbl):
    """
    Returns one of: 'key_question', 'phase', 'diff_standard', 'diff_ibl',
                    'ibl_component', 'unknown'
    """
    if not tbl.rows: return 'unknown'
    n_rows = len(tbl.rows)
    n_cols = len(tbl.columns)
    first_cell = clean(tbl.rows[0].cells[0].text)

    if n_rows == 1 and n_cols == 1 and first_cell.lower().startswith('key question'):
        return 'key_question'

    headers = [clean(c.text).lower() for c in tbl.rows[0].cells]

    if n_cols == 2 and headers[0] in ('component',''):
        return 'ibl_component'

    has_teacher = any('teacher' in h for h in headers)
    has_student = any('student' in h for h in headers)
    if has_teacher and has_student:
        return 'phase'

    has_lower  = any('lower'  in h for h in headers)
    has_higher = any('higher' in h for h in headers)
    if has_lower and has_higher:
        if n_cols >= 6: return 'diff_ibl'
        return 'diff_standard'

    # IBL diff matrix: header row may have 'more' repeated
    if n_cols >= 6 and any('more' in h or 'essential' in h for h in headers):
        return 'diff_ibl'

    return 'unknown'

# ── phase table ────────────────────────────────────────────────────────────
def parse_phase_table(tbl, abstract, num_to_abstract):
    rows = tbl.rows
    if not rows: return None
    col_headers = [cell_plain(c) for c in rows[0].cells]
    n_cols = len(col_headers)
    teacher_col = next((i for i,h in enumerate(col_headers) if 'teacher' in h.lower()), 1)
    student_col = next((i for i,h in enumerate(col_headers) if 'student' in h.lower()), 2)
    aims_col    = next((i for i,h in enumerate(col_headers)
                        if any(k in h.lower() for k in ['align','aim','relev'])),
                       n_cols-1 if n_cols>3 else None)
    parsed_rows = []
    for row in rows[1:]:
        cells = row.cells
        if not cells or not any(cell_plain(c) for c in cells): continue
        phase_raw = cell_plain(cells[0])
        phase_label = re.sub(r'\s+',' ', phase_raw).strip()
        phase_label = re.sub(r'^\d+\.\s*','', phase_label)
        teacher_bullets = cell_bullets(cells[teacher_col],abstract,num_to_abstract) if teacher_col<len(cells) else []
        student_bullets = cell_bullets(cells[student_col],abstract,num_to_abstract) if student_col<len(cells) else []
        aims_blocks = parse_aims_cell(cells[aims_col]) if aims_col is not None and aims_col<len(cells) else None
        if phase_label and not teacher_bullets and not student_bullets and not aims_blocks: continue
        parsed_rows.append({'phase':phase_label,'teacher':teacher_bullets,'student':student_bullets,'aims':aims_blocks})
    return {'col_headers':col_headers,'teacher_col':teacher_col,'student_col':student_col,'aims_col':aims_col,'rows':parsed_rows}

def parse_diff_table_standard(tbl):
    rows = tbl.rows
    if not rows: return {'type':'standard','headers':[],'stages':[],'has_role_col':False}
    headers = [cell_plain(c) for c in rows[0].cells]
    lower_col  = next((i for i,h in enumerate(headers) if 'lower'  in h.lower()),1)
    higher_col = next((i for i,h in enumerate(headers) if 'higher' in h.lower()),2)
    role_col   = next((i for i,h in enumerate(headers) if h.lower()=='role'),None)
    stages = []; UNNAMED = '__UNNAMED__'
    for row in rows[1:]:
        cells = row.cells
        if not cells or not any(cell_plain(c) for c in cells): continue
        stage_val  = cell_plain(cells[0])
        lower_val  = cell_plain(cells[lower_col])  if lower_col <len(cells) else ''
        higher_val = cell_plain(cells[higher_col]) if higher_col<len(cells) else ''
        role_val   = cell_plain(cells[role_col])   if role_col is not None and role_col<len(cells) else None
        if role_col is not None:
            if stage_val and (not stages or stage_val!=stages[-1]['stage']):
                stages.append({'stage':stage_val,'teacher_lower':'','teacher_higher':'','student_lower':'','student_higher':''})
            if stages:
                if role_val and 'teacher' in role_val.lower():
                    stages[-1]['teacher_lower']=lower_val; stages[-1]['teacher_higher']=higher_val
                elif role_val and 'student' in role_val.lower():
                    stages[-1]['student_lower']=lower_val; stages[-1]['student_higher']=higher_val
        else:
            if stage_val:
                stages.append({'stage':stage_val,'lower':lower_val,'higher':higher_val})
            else:
                if not stages: stages.append({'stage':UNNAMED,'lower':lower_val,'higher':higher_val})
                else:
                    stages[-1]['lower']=(stages[-1].get('lower','')+' '+lower_val).strip()
                    stages[-1]['higher']=(stages[-1].get('higher','')+' '+higher_val).strip()
    return {'type':'standard','headers':headers,'stages':stages,'has_role_col':role_col is not None,'unnamed_sentinel':UNNAMED}

def parse_diff_table_ibl(tbl):
    rows = tbl.rows
    if not rows: return {'type':'ibl_matrix','features':[]}
    features = []; seen = {}
    for row in rows[2:]:
        cells = row.cells
        if len(cells)<6: continue
        feature_val=clean(cells[0].text); role_val=clean(cells[1].text)
        level_cells=[clean(cells[i].text) for i in range(2,6)]
        if not feature_val and not role_val: continue
        if feature_val and feature_val not in seen:
            seen[feature_val]=len(features)
            features.append({'name':feature_val,'teacher':[],'students':[]})
        fi=seen.get(feature_val,len(features)-1)
        if fi>=len(features): continue
        if 'teacher' in role_val.lower(): features[fi]['teacher']=level_cells
        elif 'student' in role_val.lower(): features[fi]['students']=level_cells
    return {'type':'ibl_matrix','features':features}

# ── free-flowing section stream ────────────────────────────────────────────
def stream_section(body, doc, docx_path, abstract, num_to_abstract,
                   start_marker_fn, end_marker_fn,
                   table_handler=None):
    """
    Walk document body from start_marker to end_marker, emitting content items:
      {type:'bullet', text, tech, level}
      {type:'image',  src, caption}
      {type:'table',  table_type, data}   -- if table_handler provided
      {type:'paragraph', text}            -- plain non-bullet paragraph
    """
    items = []; in_section = False; last_image_idx = None

    for child in body:
        tag = child.tag.split('}')[-1]

        if tag == 'p':
            from docx.text.paragraph import Paragraph as DP
            p = DP(child, doc)
            text  = para_text(p)
            style = para_style(p)
            tl    = text.lower()

            if not in_section:
                if start_marker_fn(tl, style):
                    in_section = True
                continue

            # end of section
            if end_marker_fn(tl, style):
                break

            # image
            if has_image(p):
                src = extract_image_b64(p, docx_path)
                if src:
                    caption = re.sub(r'^\.\s*','', text).strip()
                    items.append({'type':'image','src':src,'caption':caption})
                    last_image_idx = len(items)-1
                continue

            # caption style — attach to preceding image
            if 'Caption' in style and last_image_idx is not None:
                if not items[last_image_idx]['caption']:
                    items[last_image_idx]['caption'] = re.sub(r'^\.\s*','',text).strip()
                continue

            if not text: continue

            # bullet or plain paragraph
            level = get_bullet_level(p, abstract, num_to_abstract)
            tech  = is_tech(p)
            items.append({'type':'bullet','text':strip_tech(text) if tech else text,
                          'tech':tech,'level':level})

        elif tag == 'tbl' and in_section:
            from docx.table import Table as DT
            tbl = DT(child, doc)
            ttype = identify_table(tbl)

            if table_handler:
                result = table_handler(tbl, ttype)
                if result is not None:
                    items.append(result)
            # else: skip tables (e.g. in LE preamble scanning)

    return items

# ── section marker helpers ─────────────────────────────────────────────────
WHAT_IS_START = lambda tl,st: 'what is' in tl and any(s in st for s in ('Title','Heading','Normal'))
LE_START      = lambda tl,st: any(m in tl for m in ['le facilitated','learning experience (le) in the classroom'])
DIFF_START    = lambda tl,st: 'differentiating' in tl and any(s in st for s in ('Heading','Title','Subtitle','heading'))

def make_end(*markers):
    return lambda tl,st: any(m in tl for m in markers) and ('Heading' in st or 'Title' in st)

WHAT_IS_END = make_end('le facilitated','learning experience (le) in the classroom','differentiating')
LE_END      = make_end('differentiating')
DIFF_END    = make_end('what is')  # safety — shouldn't appear after diff

# ── key question (from table) ──────────────────────────────────────────────
def parse_key_question_paragraphs(doc):
    """Return list of paragraph strings from the Key Question table cell."""
    for tbl in doc.tables:
        if tbl.rows and tbl.columns:
            cell = tbl.rows[0].cells[0]
            text = clean(cell.text)
            if text.lower().startswith('key question'):
                paras = [clean(p.text) for p in cell.paragraphs if clean(p.text)]
                return paras if paras else [text]
    return None

# ── per-doc orchestrator ───────────────────────────────────────────────────
def parse_doc(docx_path, approach_key):
    doc = Document(docx_path)
    abstract, num_to_abstract = build_numbering_lookup(doc)
    body = doc.element.body
    data = {}

    # ── What is X? — fully free-flowing ──
    def what_is_table_handler(tbl, ttype):
        if ttype == 'ibl_component':
            # flatten to bullets
            rows = tbl.rows
            bullets = []
            for row in rows[1:]:
                num = clean(row.cells[0].text); act = clean(row.cells[1].text)
                if num and act:
                    bullets.append({'type':'bullet','text':f'{num}. {act}','tech':False,'level':0})
            return {'type':'bullet_group','items':bullets}
        return None  # skip other tables in what_is

    what_is_items = stream_section(
        body, doc, docx_path, abstract, num_to_abstract,
        WHAT_IS_START, WHAT_IS_END,
        table_handler=what_is_table_handler
    )
    data['what_is'] = what_is_items

    # ── LE facilitated — free-flowing preamble, then structured phase table ──
    # Walk LE section: collect preamble items until we hit the phase table
    le_preamble = []
    phase_table_data = None
    key_question_paras = parse_key_question_paragraphs(doc)
    if key_question_paras:
        data['key_question_paras'] = key_question_paras

    in_le = False
    found_phase = False
    for child in body:
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            from docx.text.paragraph import Paragraph as DP
            p = DP(child, doc)
            text = para_text(p); style = para_style(p); tl = text.lower()
            if not in_le:
                if LE_START(tl, style): in_le = True
                continue
            if LE_END(tl, style): break
            if found_phase: continue  # after phase table, skip remaining LE paras
            if not text and not has_image(p): continue
            if has_image(p):
                src = extract_image_b64(p, docx_path)
                if src:
                    caption = re.sub(r'^\.\s*','',text).strip()
                    le_preamble.append({'type':'image','src':src,'caption':caption})
                continue
            level = get_bullet_level(p, abstract, num_to_abstract)
            tech  = is_tech(p)
            le_preamble.append({'type':'bullet','text':strip_tech(text) if tech else text,
                                 'tech':tech,'level':level})
        elif tag == 'tbl' and in_le:
            from docx.table import Table as DT
            tbl = DT(child, doc)
            ttype = identify_table(tbl)
            if ttype == 'key_question':
                continue  # already handled separately
            if ttype == 'phase' and not found_phase:
                phase_table_data = parse_phase_table(tbl, abstract, num_to_abstract)
                found_phase = True

    data['le_preamble']  = le_preamble
    if phase_table_data:
        data['phase_table'] = phase_table_data

    # ── Differentiating Instruction — fully free-flowing ──
    def diff_table_handler(tbl, ttype):
        if ttype == 'diff_ibl':
            return {'type':'table','table_type':'diff_ibl','data':parse_diff_table_ibl(tbl)}
        if ttype == 'diff_standard':
            return {'type':'table','table_type':'diff_standard','data':parse_diff_table_standard(tbl)}
        return None

    diff_items = stream_section(
        body, doc, docx_path, abstract, num_to_abstract,
        DIFF_START, DIFF_END,
        table_handler=diff_table_handler
    )
    data['diff_stream'] = diff_items

    return data

# ── main ───────────────────────────────────────────────────────────────────
def main():
    docs = {
        'pbl': SCRIPT_DIR/'PBL.docx',
        'ibl': SCRIPT_DIR/'IBL.docx',
        'cbl': SCRIPT_DIR/'CBL.docx',
        'ssi': SCRIPT_DIR/'SSI.docx',
    }
    approach_data = {}
    for key, path in docs.items():
        if not path.exists():
            print(f'WARNING: {path} not found'); approach_data[key]={}; continue
        print(f'Parsing {path.name}...')
        approach_data[key] = parse_doc(str(path), key)
        print(f'  Done: {key}')
    template_path = SCRIPT_DIR/'template.html'
    output_path   = SCRIPT_DIR/'pedagogy-hub.html'
    if not template_path.exists():
        print(f'ERROR: template.html not found'); return
    template = template_path.read_text(encoding='utf-8')
    for key, data in approach_data.items():
        template = template.replace(f'__{key.upper()}_DATA__', json.dumps(data, ensure_ascii=False))
    output_path.write_text(template, encoding='utf-8')
    print(f'\nOutput written to {output_path}')

if __name__ == '__main__':
    main()
